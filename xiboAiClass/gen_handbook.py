#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ===== 页面设置 =====
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = section.right_margin = Cm(2.5)
section.top_margin  = section.bottom_margin = Cm(2.2)

# ===== 样式助手 =====
def set_font(run, name='PingFang SC', size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '微软雅黑')
    rPr.insert(0, rFonts)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level==1 else 10)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    sizes = {1: 20, 2: 15, 3: 13}
    bolds = {1: True, 2: True, 3: True}
    colors= {1: (17,17,17), 2: (40,40,40), 3: (60,60,60)}
    set_font(run, size=sizes.get(level,12), bold=bolds.get(level,False), color=colors.get(level))
    return p

def add_body(doc, text, indent=False, color=(80,80,80)):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    set_font(run, size=10.5, color=color)
    return p

def add_tip(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run('▸  ' + text)
    set_font(run, size=10, color=(100,100,100))

def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.6)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(30, 30, 30)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'F5F5F5')
    p._p.get_or_add_pPr().append(shading)

def add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ===========================
# 封面
# ===========================
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(60)
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('喜播有声 · AI 内部特训营')
set_font(run, size=9, color=(150,150,150))

p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(16)
run = p.add_run('让 AI 替你打工')
set_font(run, size=32, bold=True, color=(17,17,17))

p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Claude 两日小白特训营 · 学员手册')
set_font(run, size=14, color=(100,100,100))

doc.add_paragraph()
p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('理论 30% · 实操 70% · 训战结合')
set_font(run, size=10, color=(150,150,150))

p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('主讲：蒋德铭  |  特邀：周鸿斌（喜马拉雅珠峰 AI 负责人）')
set_font(run, size=10, color=(150,150,150))

doc.add_page_break()

# ===========================
# 使用说明
# ===========================
add_heading(doc, '📖 这本手册怎么用', 1)
add_body(doc, '这不是一本读完就放在一边的资料，而是你在培训现场的操作指南。')
add_body(doc, '每个实操环节，手册里都有：')
add_tip(doc, '步骤说明：告诉你下一步做什么')
add_tip(doc, '提示词模板：可以直接复制粘贴给 Claude')
add_tip(doc, '检查清单：帮你判断自己做得对不对')
add_body(doc, '培训结束后，带着这本手册回去，可以在自己部门重新演练一遍。')
add_hr(doc)

# ===========================
# 第一天
# ===========================
add_heading(doc, '第一天｜核心理论 · 先换大脑再动手', 1)
add_body(doc, '时长约 2 小时。每个模块先看演示，再自己动手。', color=(120,120,120))

# 模块1
add_heading(doc, '模块 1｜思维大换血', 2)
add_heading(doc, '核心观点', 3)
add_body(doc, '别把 AI 当打字机——要当成不用睡觉的员工团队。')
add_body(doc, '你只需要做两件事：')
add_tip(doc, '提需求：告诉 AI 要做什么（不是怎么做）')
add_tip(doc, '微决策：在 AI 给的 2-3 个方案里，选一个或稍作调整')
add_heading(doc, '你现在的思考题', 3)
add_body(doc, '在下方写下：我工作中每周重复 2 次以上的事情是什么？')
add_body(doc, '（今天下午实战会用到这个答案）')
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
for _ in range(3):
    p = doc.add_paragraph('_' * 60)
    p.paragraph_format.space_after = Pt(10)
    run = p.runs[0]
    run.font.color.rgb = RGBColor(200,200,200)
    run.font.size = Pt(10)
add_hr(doc)

# 模块2
add_heading(doc, '模块 2｜Project 专属记忆库', 2)
add_heading(doc, '解决的问题', 3)
add_body(doc, '每次新对话，AI 都失忆了——你得从头解释公司背景、写作风格、禁止用词。')
add_body(doc, 'Project 让 AI 永久记住你的背景，以后直接说"帮我写通知"，它自带公司语气出稿。')
add_heading(doc, '操作步骤', 3)
add_tip(doc, '1. 打开 Claude 桌面版 → 点击左侧「项目」→「新建项目」')
add_tip(doc, '2. 上传文件：公司简介、部门历史文件、过去写得好的模板')
add_tip(doc, '3. 在「项目说明」里填写以下内容（参考下方模板）')
add_tip(doc, '4. 保存，以后每次在项目里对话，AI 都带着这些记忆')
add_heading(doc, '项目说明填写模板（复制后替换括号内容）', 3)
add_code(doc, '你是（你的公司/部门名称）的（你的角色，如：行政助理/内容编辑）。')
add_code(doc, '写作风格：（正式/亲切/简洁——选一个或自己描述）')
add_code(doc, '禁止用词：赋能、闭环、链路（可以加你自己不喜欢的词）')
add_code(doc, '公司全称：（填写公司完整名称）')
add_code(doc, '过去写得好的文件风格可以参考我上传的附件。')
add_heading(doc, '动手任务', 3)
add_tip(doc, '新建一个以你部门为名的 Project')
add_tip(doc, '上传至少 1 份部门真实文件')
add_tip(doc, '填好项目说明')
add_tip(doc, '测试：说一件你日常工作里的事，让 AI 帮你做，看输出质量')
add_hr(doc)

# 模块3
add_heading(doc, '模块 3｜Cowork · 定点手术式修改', 2)
add_heading(doc, '解决的问题', 3)
add_body(doc, '改一个地方，AI 把全文重写——这叫"刷屏折磨"。')
add_body(doc, 'Cowork 让你用鼠标框选要改的那一段，AI 只改那里，其他不动。')
add_heading(doc, '操作方法', 3)
add_tip(doc, '1. 在 Cowork 区域生成完整草稿')
add_tip(doc, '2. 鼠标框选你想修改的段落')
add_tip(doc, '3. 在弹出框里说你的修改要求')
add_tip(doc, '4. AI 只修改选中部分，完成后继续框选下一处')
add_heading(doc, '练习提示词', 3)
add_code(doc, '帮我生成一份商场亲子游园活动策划，')
add_code(doc, '包括：活动主题、时间安排、费用预算三个部分。')
add_body(doc, '生成后，尝试完成 3 次框选修改：修改主题 / 修改语气 / 自由修改一处')
add_hr(doc)

# 模块4
add_heading(doc, '模块 4｜Skill · 标准化流水线', 2)
add_heading(doc, '解决的问题', 3)
add_body(doc, '重复工作每次都要重新说规矩，效率低、质量不稳定。')
add_body(doc, 'Skill 是你写给 AI 的"岗位说明书"，定好规矩，以后只需扔进原材料，自动出成品。')
add_heading(doc, 'Skill = 说明书 + 模板 + 检查清单', 3)
add_heading(doc, '三步创建你的第一个 Skill', 3)
add_tip(doc, '第一步：写"标准参考文件"——列出 5-8 条这件事的做法规范')
add_tip(doc, '第二步：把规范粘贴给 Claude，让它帮你生成 Skill 文档')
add_tip(doc, '第三步：用真实素材测试，发现问题告诉 Claude 修改')
add_heading(doc, '创建 Skill 的提示词', 3)
add_code(doc, '我上传了一份关于（你的工作内容）的标准规范文件。')
add_code(doc, '请根据这份文件，帮我创建一个专业的 Skill，包含：')
add_code(doc, '1. 角色设定  2. 工作流程  3. 风格规范')
add_code(doc, '4. 输出格式模板  5. 质量检查清单')
add_code(doc, '输出为完整 Skill 文档，格式清晰，可以直接保存使用。')

doc.add_page_break()

# ===========================
# 第二天
# ===========================
add_heading(doc, '第二天｜训战实操工坊（约半天）', 1)
add_body(doc, '带着昨天写下的"每周重复最多的那件事"来，今天把它变成一条自动流水线。', color=(120,120,120))

# 关卡01
add_heading(doc, '🔴 关卡 01｜喂背景，立规矩（45 分钟）', 2)
add_body(doc, '标签：企业经营管理  |  目标：产出你的专属 Project + 个人 Skill')
add_heading(doc, '步骤 1：完善 Project 记忆库', 3)
add_tip(doc, '✅ 上传部门历史文件（至少 1 份）')
add_tip(doc, '✅ 上传你过去写得好的文案作为风格参考')
add_tip(doc, '✅ 完善项目说明（角色 / 风格 / 禁止词）')
add_heading(doc, '步骤 2：创建你的个人 Skill', 3)
add_tip(doc, '1. 回想你工作里重复最多的那件事')
add_tip(doc, '2. 写出 5-8 条"标准做法"')
add_tip(doc, '3. 让 Claude 帮你生成 Skill')
add_tip(doc, '4. 测试，找问题，反馈修改')
add_heading(doc, '步骤 3：实战验证', 3)
add_body(doc, '用你的 Project + Skill，完成一件真实工作任务。')
add_tip(doc, '参考：起草本月部门工作汇报 / 写一封内部通知 / 整理会议纪要摘要')
add_heading(doc, '成果检查清单', 3)
add_tip(doc, '□ 输出内容有没有公司语气？')
add_tip(doc, '□ 风格和你过去写的文件像不像？')
add_tip(doc, '□ 有没有出现你不希望出现的词？')
add_hr(doc)

# 关卡02
add_heading(doc, '🔴 关卡 02｜找痛点，共修改（60 分钟）', 2)
add_body(doc, '标签：商业体运营策划  |  目标：用 Cowork 完成 3 次精准框选修改')
add_heading(doc, '主任务提示词', 3)
add_code(doc, '帮我生成一份完整的商场年中大促招商计划书，包含：')
add_code(doc, '1. 活动背景与目标')
add_code(doc, '2. 活动时间与主题')
add_code(doc, '3. 招商对象与条件')
add_code(doc, '4. 商家权益说明')
add_code(doc, '5. 申请流程与截止时间')
add_code(doc, '商场名称：[填写你的商场名称]')
add_heading(doc, '3 次框选修改任务', 3)
add_tip(doc, '第 1 次：框选"活动主题" → 改成夏季清凉概念，要适合做横幅标语')
add_tip(doc, '第 2 次：框选"商家权益" → 语气改成更吸引商家、突出实际好处')
add_tip(doc, '第 3 次：自己找一处不满意的地方，框选后修改')
add_hr(doc)

# 关卡03
add_heading(doc, '🔴 关卡 03｜定流程，开工厂（60 分钟）', 2)
add_body(doc, '标签：有声作品量制作  |  目标：建立「去冗白→标台词→加音效」三步流水线')
add_heading(doc, '有声脚本处理规范（用于创建 Skill）', 3)
add_code(doc, '【有声脚本标准化处理规范】')
add_code(doc, '')
add_code(doc, '第一步（去冗白）：')
add_code(doc, '- 删除过度文学化描写，保留情节推进必要的描述')
add_code(doc, '- 长句拆短，每句不超过25字')
add_code(doc, '- 对话一字不改')
add_code(doc, '')
add_code(doc, '第二步（标台词）：')
add_code(doc, '- 旁白标注：【旁白】')
add_code(doc, '- 角色对话标注：【角色名】（情绪）')
add_code(doc, '- 情绪：（轻声）（激动）（冷淡）等')
add_code(doc, '')
add_code(doc, '第三步（加音效）：')
add_code(doc, '- 场景切换：[场景音效-XX]')
add_code(doc, '- 停顿：[停顿1秒] [停顿2秒]')
add_code(doc, '- 动作：[动作音效-XX]')
add_heading(doc, '练习素材（可直接粘贴给 Claude 测试）', 3)
add_code(doc, '天色将晚，林中的风凉了下来。苏晚站在溪边，看着水里自己的倒影，')
add_code(doc, '心里说不清楚是什么滋味。"你还在这里？"身后有人说。')
add_code(doc, '她没有回头，知道是谁。"你来了。""嗯。"沈屿走到她身边，站住，')
add_code(doc, '也看着水里。"今天怎么样？""没怎么样，"她说，"就这样。"')
add_heading(doc, '成果检查清单', 3)
add_tip(doc, '□ 有没有去掉过于文学化的描写？')
add_tip(doc, '□ 每个角色的对话都有标注吗？')
add_tip(doc, '□ 音效标注位置合理吗？')
add_tip(doc, '□ 配音员拿到这份脚本，能直接录音吗？')

doc.add_page_break()

# ===========================
# 学完带走
# ===========================
add_heading(doc, '🎯 学完带走的六件东西', 1)
items = [
    ('老板思维', '从自己用 AI 干活，变成让 AI 替你干活'),
    ('专属记忆库', '一个装了公司背景和你风格偏好的 Project'),
    ('个人技能包', '至少一个针对你工作场景的 Skill'),
    ('协作技巧', 'Cowork 框选精准修改，不再刷屏重写'),
    ('每天省 2 小时', '流水线替代重复工作，时间还给重要的事'),
    ('三份真实业务文件', '两天培训中实际产出的真实工作成果'),
]
for i, (name, desc) in enumerate(items, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run1 = p.add_run(f'{i}. {name}  ')
    set_font(run1, size=11, bold=True, color=(17,17,17))
    run2 = p.add_run(desc)
    set_font(run2, size=10.5, color=(100,100,100))
add_hr(doc)

# ===========================
# 常用提示词速查
# ===========================
add_heading(doc, '⚡ 常用提示词速查表', 1)

prompts = [
    ('创建 Project 项目说明', '你是（公司/部门）的（角色）。\n写作风格：（正式/简洁）。禁止用词：赋能、闭环。\n公司全称：（全称）。'),
    ('让 AI 起草通知', '基于我上传的文件背景，\n帮我起草一份（通知主题）的内部通知，\n语气（正式/亲切），不超过 300 字。'),
    ('创建 Skill 指令', '根据我上传的规范文件，\n帮我创建一个完整的 Skill，包含：\n角色设定、工作流程、风格规范、输出格式、质量清单。'),
    ('测试并修改 Skill', '用这个 Skill 处理以下内容：\n（粘贴原始素材）\n\n发现问题后说：\n把 Skill 里的（某条规范）改一下，\n要求（新的要求）。'),
    ('有声脚本处理', '调用有声脚本 Skill，\n处理以下小说原文，\n输出可直接录音的脚本格式：\n（粘贴小说原文）'),
]
for title, prompt in prompts:
    add_heading(doc, title, 3)
    add_code(doc, prompt)
    doc.add_paragraph()

# ===========================
# 页脚信息
# ===========================
add_hr(doc)
p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('喜播有声 AI 内部特训营 · 学员手册 · 2026')
set_font(run, size=9, color=(180,180,180))

# ===== 保存 =====
output = '/Users/david/xibo/xiboweb/xiboAiClass/学员手册_喜播AI特训营.docx'
doc.save(output)
print(f'✅ 手册已生成：{output}')
